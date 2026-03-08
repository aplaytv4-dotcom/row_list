import os
import sys
from collections import defaultdict
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


FONT_NAME = "Times New Roman"
FONT_SIZE = 14


def get_app_dir() -> str:
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def get_reports_dir(report_date=None) -> str:
    if report_date is None:
        report_date = datetime.now()

    app_dir = get_app_dir()
    month_folder = report_date.strftime("%Y-%m")
    reports_dir = os.path.join(app_dir, "отчеты", month_folder)
    os.makedirs(reports_dir, exist_ok=True)
    return reports_dir


def safe_filename(text: str) -> str:
    bad = '<>:"/\\|?*'
    for ch in bad:
        text = text.replace(ch, "_")
    return text


def open_file(filepath: str):
    try:
        os.startfile(filepath)
    except Exception:
        pass


def save_document(doc: Document, filename: str, report_date=None) -> str:
    filename = safe_filename(filename)
    reports_dir = get_reports_dir(report_date)
    path = os.path.join(reports_dir, filename)
    doc.save(path)
    open_file(path)
    return path


def set_run_font(run, bold=False, size=FONT_SIZE):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold


def ensure_doc_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(FONT_SIZE)


def add_single_line(doc: Document, text="", bold=False, align=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.25

    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold)

    if align is not None:
        p.alignment = align

    return p


def add_block_lines(doc: Document, lines, bold_lines=None, align=None):
    """
    Все строки пишутся внутри одного абзаца через run + '\\n'
    """
    if bold_lines is None:
        bold_lines = set()

    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.25

    if align is not None:
        p.alignment = align

    total = len(lines)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, bold=(idx in bold_lines))
        if idx != total - 1:
            run.add_break()

    return p


def add_signature_block(doc: Document, chief_name: str, current_date: str):
    add_single_line(doc, "")

    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1

    left = p.add_run("Начальник отдела")
    set_run_font(left, bold=True)

    if chief_name.strip():
        spacer = p.add_run(" " * 25)
        set_run_font(spacer)
        right = p.add_run(chief_name.strip())
        set_run_font(right, bold=True)

    #add_single_line(doc, "")
    add_single_line(doc, f"Дата составления: {current_date}")


def plural_people(n: int) -> str:
    n = abs(int(n))
    if 11 <= (n % 100) <= 14:
        return "человек"
    if n % 10 == 1:
        return "человек"
    if n % 10 in (2, 3, 4):
        return "человека"
    return "человек"


def plural_case(n: int) -> str:
    n = abs(int(n))
    if 11 <= (n % 100) <= 14:
        return "случаев"
    if n % 10 == 1:
        return "случай"
    if n % 10 in (2, 3, 4):
        return "случая"
    return "случаев"


def format_people(n: int) -> str:
    return f"{n} {plural_people(n)}"


def format_date_ru(date_str: str) -> str:
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d.%m.%Y")
    except Exception:
        return date_str


def build_daily_reason_groups(absent_rows, date_view: str):
    groups = defaultdict(list)

    for row in absent_rows:
        reason = (row.get("reason") or "").strip()
        comment = (row.get("comment") or "").strip()
        name = row.get("name", "").strip()

        if not reason:
            reason = "Не указано"

        if reason == "Уважительная причина":
            line = f"{name} — {comment if comment else '-'} — {date_view}"
        else:
            line = f"{name} — {date_view}"

        groups[reason].append(line)

    return groups


def generate_daily_report(
    daily_data,
    staff_total_with_vacancies: int,
    chief_name: str,
    report_date_db: str | None = None,
):
    if report_date_db:
        report_dt = datetime.strptime(report_date_db, "%Y-%m-%d")
    else:
        report_dt = datetime.now()

    today_db = report_dt.strftime("%Y-%m-%d")
    today_view = report_dt.strftime("%d.%m.%Y")

    doc = Document()
    ensure_doc_style(doc)

    add_single_line(doc, "ОТЧЁТ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_single_line(doc, "о присутствии сотрудников", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_single_line(doc, "")
    add_single_line(doc, f"Дата: {today_view}", align=WD_ALIGN_PARAGRAPH.RIGHT)

    real_staff = len(daily_data)
    present_rows = [row for row in daily_data if row.get("status") == "Присутствует"]
    absent_rows = [row for row in daily_data if row.get("status") != "Присутствует"]

    present_count = len(present_rows)
    absent_count = len(absent_rows)

    head_lines = [
        "",
        "По штату:",
        f"С учётом вакантных должностей: {format_people(staff_total_with_vacancies)}",
        f"Без учёта вакантных должностей: {format_people(real_staff)}",
        "",
        f"Присутствуют: {format_people(present_count)}",
        f"Отсутствуют: {format_people(absent_count)}",
    ]
    add_block_lines(doc, head_lines, bold_lines={1, 5, 6})

    reason_groups = build_daily_reason_groups(absent_rows, today_view)

    preferred_order = [
        "Уважительная причина",
        "Больничный",
        "Отпуск",
        "Командировка",
        "Дежурство",
        "Отпросился",
        "Не указано",
    ]

    titles = {
        "Уважительная причина": "Отпросились по уважительной причине",
        "Больничный": "Больничный",
        "Отпуск": "Отпуск",
        "Командировка": "Командировка",
        "Дежурство": "Дежурство",
        "Отпросился": "Отпросился",
        "Не указано": "Причина не указана",
    }

    has_any_group = any(len(lines) > 0 for lines in reason_groups.values())

    if has_any_group:
        #add_single_line(doc, "")
        add_single_line(doc, "Из них:")
        used = set()

        for reason in preferred_order:
            if reason in reason_groups and reason_groups[reason]:
                lines = reason_groups[reason]
                block = [f"{titles.get(reason, reason)}: {format_people(len(lines))}"]
                for i, line in enumerate(lines, start=1):
                    block.append(f"{i}. {line}")
                add_block_lines(doc, block)
                used.add(reason)

        for reason, lines in reason_groups.items():
            if reason not in used and lines:
                block = ["", f"{reason}: {format_people(len(lines))}"]
                for i, line in enumerate(lines, start=1):
                    block.append(f"{i}. {line}")
                add_block_lines(doc, block)
    else:
        #add_single_line(doc, "")
        add_single_line(doc, "Отсутствующие сотрудники не зафиксированы.")

    add_signature_block(doc, chief_name, today_view)

    filename = f"attendance_daily_{today_db}.docx"
    return save_document(doc, filename, report_dt)


def generate_period_report(
    title: str,
    date_from: str,
    date_to: str,
    employees,
    attendance_rows,
    staff_total_with_vacancies: int,
    chief_name: str,
):
    now = datetime.now()
    current_date = now.strftime("%d.%m.%Y")
    date_from_view = format_date_ru(date_from)
    date_to_view = format_date_ru(date_to)

    doc = Document()
    ensure_doc_style(doc)

    staff_total_real = len(employees)
    emp_names = {
        emp_id: name for emp_id, department, position, rank, name in employees
    }

    absent_reason_counter = defaultdict(int)
    absent_detail_lines = []
    absent_count = 0

    for emp_id, att_date, status, reason, comment in attendance_rows:
        if status == "Присутствует":
            continue

        absent_count += 1

        reason = (reason or "").strip()
        comment = (comment or "").strip()
        emp_name = emp_names.get(emp_id, "Неизвестный сотрудник")
        date_view = format_date_ru(att_date)

        if reason:
            absent_reason_counter[reason] += 1
        else:
            reason = "Не указано"
            absent_reason_counter[reason] += 1

        if reason == "Уважительная причина":
            detail_text = f"{emp_name} — {reason.lower()}, {comment if comment else '-'} — {date_view}"
        elif reason == "Не указано":
            detail_text = f"{emp_name} — причина не указана — {date_view}"
        else:
            detail_text = f"{emp_name} — {reason.lower()} — {date_view}"

        absent_detail_lines.append(detail_text)

    add_single_line(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    main_lines = [
        "",
        "Отчёт за период:",
        f"{date_from_view} — {date_to_view}",
        "",
        "По штату:",
        f"С учётом вакантных должностей: {staff_total_with_vacancies} человек",
        f"Без учёта вакантных должностей: {staff_total_real} человек",
        "",
        f"Отсутствовали за период: {absent_count} {plural_case(absent_count)}",
        "Из них:",
    ]
    add_block_lines(doc, main_lines, bold_lines={4, 8})

    preferred_order = [
        "Уважительная причина",
        "Больничный",
        "Отпуск",
        "Командировка",
        "Дежурство",
        "Отпросился",
        "Не указано",
    ]

    reason_lines = []
    used = set()

    for reason_name in preferred_order:
        if reason_name in absent_reason_counter:
            count = absent_reason_counter[reason_name]
            reason_lines.append(f"{reason_name.lower()}: {count} {plural_case(count)}")
            used.add(reason_name)

    for reason_name, count in absent_reason_counter.items():
        if reason_name not in used:
            reason_lines.append(f"{reason_name.lower()}: {count} {plural_case(count)}")

    if reason_lines:
        add_block_lines(doc, reason_lines)
    else:
        add_single_line(doc, "- отсутствий не зафиксировано")

    add_single_line(doc, "")

    if absent_detail_lines:
        detail_lines = [f"{i}. {line}" for i, line in enumerate(absent_detail_lines, start=1)]
        add_block_lines(doc, detail_lines)
    else:
        add_single_line(doc, "Случаи отсутствия за период не зафиксированы.")

    add_signature_block(doc, chief_name, current_date)

    filename = f"{title.lower().replace(' ', '_')}_{date_from}_to_{date_to}.docx"
    return save_document(doc, filename, now)